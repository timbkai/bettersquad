# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)

def h(text, level):
    p = doc.add_heading(text, level=level)
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullet(label, rest):
    p = doc.add_paragraph(style='List Bullet')
    if label:
        r = p.add_run(label + ': ')
        r.bold = True
    p.add_run(rest)
    return p

def callout(title, body):
    """Shaded one-row info box."""
    t = doc.add_table(rows=2, cols=1)
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0 = t.rows[0].cells[0].paragraphs[0]
    r = c0.add_run(title)
    r.bold = True
    t.rows[1].cells[0].paragraphs[0].add_run(body)
    doc.add_paragraph()
    return t

# ---------------- Title block ----------------
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('BetterSquad  ·  Businessplan')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = ACCENT

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Investiere in deinen Kader. Nicht in einzelne Spieler:innen.')
r.italic = True
r.font.size = Pt(11)

mp = doc.add_paragraph()
mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = mp.add_run('WiSo Köln · Bachelorseminar · Wintersemester 2024/25')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()

# ---------------- Heading 1 ----------------
h('3  Produkt / Dienstleistung', 1)
para('Dieses Kapitel beschreibt das Leistungsangebot von BetterSquad entlang der '
     'beiden untrennbaren Bausteine Hardware und Software. BetterSquad ist ein '
     'integriertes Belastungssteuerungs-System für den Amateur- und Semi-Pro-'
     'Frauenfußball im DACH-Raum, das professionelle Methodik der '
     'Belastungssteuerung für Frauenabteilungen verfügbar macht – als '
     'schlüsselfertiges Komplettsystem aus tragbarer Sensorik und cloudbasierter '
     'Analysesoftware.')

# ============================================================
# 3.1  Beschreibung des Produktes / der Dienstleistung
# ============================================================
h('3.1  Beschreibung des Produktes / der Dienstleistung', 2)

para('BetterSquad verbindet vereinseigene Sensor-Hardware im Pool-Modell mit einer '
     'Software-Plattform, die rohe Mess- und Befindlichkeitsdaten nicht als '
     'Performance-Dashboard ausgibt, sondern zu einer konkreten, individuell auf die '
     'Spielerin kalibrierten Trainingsempfehlung verdichtet. Der inhaltliche Kern ist '
     'damit kein Tracking-Gadget, sondern ein female-informed Decision Layer: eine '
     'Entscheidungsschicht, die die äußere Belastung (GPS), die innere Belastung '
     '(Herzfrequenz und HRV) und das subjektive Befinden (Hooper-Index) auf Basis '
     'individueller Baselines jeder Spielerin auswertet.')

para('Das Produkt besteht aus drei aufeinander abgestimmten Bausteinen: (1) einer '
     'vereinseigenen Hardware im geteilten Sensor-Pool, (2) einer Software-Plattform '
     'aus Trainer:innen-Dashboard und Spieler:innen-App sowie (3) einem '
     'female-informed Decision Layer mit datenschutzfreundlicher Architektur.')

callout('Kernversprechen',
        'Profi-Methodik der Belastungssteuerung zum Amateur-Preis: ab 219 €/Monat '
        'statt vier- bis fünfstelliger Jahresaufwände bei Pro-Gerät-Modellen – '
        'DSGVO-konform, mit Serverstandort Deutschland und female-informed '
        'Decision Layer.')

# --- 3.1.1 Hardware ---
h('3.1.1  Hardware: vereinseigener Sensor-Pool', 3)
para('Die Hardware bildet die objektive Datengrundlage des Systems. Bewusst setzt '
     'BetterSquad auf erprobte, kalibrierte Standardkomponenten statt auf teure '
     'Eigenentwicklungen. Das senkt Stückkosten und Ausfallrisiko, verkürzt die '
     'Time-to-Market und sichert eine validierte Messqualität. Die Geräte werden der '
     'Frauenabteilung als geteilter Pool zur Verfügung gestellt: Eine typische '
     'Abteilung mit drei bis fünf Teams und 60 bis 80 Spielerinnen wird so mit einem '
     'überschaubaren Gerätebestand abgedeckt.')

para('GPS-Tracker (äußere Belastung) – getragen in einer leichten Trageweste '
     'zwischen den Schulterblättern, erfasst die äußere Belastung in Echtzeit:', bold=True)
bullet('Laufleistung', 'Zurückgelegte Distanz und Distanz in Hochintensitäts-Zonen')
bullet('Sprints & Speed', 'Anzahl der Sprints sowie Höchstgeschwindigkeit (km/h)')
bullet('HI-Runs', 'High-Intensity-Runs als Indikator azyklischer Belastung')
bullet('Acc/Dec', 'Beschleunigungs- und Abbrems-Lasten zur Erkennung von Asymmetrien')
para('Aus diesen Rohdaten wird die externe Tageslast (Load) berechnet, die als '
     'Grundgröße in das ACWR-Modell einfließt.')

para('Polar-H10-Pulsgurte (innere Belastung & Recovery) – der medizinisch '
     'validierte Brustgurt misst die innere Belastung und liefert die Datenbasis '
     'für das Recovery-Monitoring:', bold=True)
bullet('HR', 'Herzfrequenz in Echtzeit zur Steuerung der Trainingsintensität')
bullet('HRV', 'Herzfrequenzvariabilität als sensibler Frühindikator der Regeneration')
bullet('Recovery', 'Trendanalyse über Tage, um Überlastung mehrere Tage früher zu erkennen')
para('Für eine preissensible Pool-Beschaffung existieren validierte Dual-Mode-'
     'Brustgurte (ANT+ und Bluetooth) wie Magene H64, Coospo H6 oder der '
     'Decathlon-KALENJI-HRM-Gurt (Einkauf rund 25–35 €); der Polar H10 bleibt die '
     'Referenz für maximale EKG-Genauigkeit und medizinische Validierung.')

para('Tracker-Dock & QR-Checkout – die Geräte lagern und laden in einem zentralen '
     'Tracker-Dock. Vor dem Training scannt jede Spielerin über die App einen '
     'QR-Code und bekommt automatisch einen freien Tracker und Pulsgurt zugewiesen:', bold=True)
bullet('Self-Checkout', 'Automatische Zuweisung von Tracker- und Gurt-Nummer per QR-Scan')
bullet('Pool-Transparenz', 'Live-Übersicht über verfügbare, ausgegebene und in Wartung befindliche Geräte')
bullet('Onboarding', 'Geführte Schritt-für-Schritt-Anlage (Weste, Gurt befeuchten, Bluetooth prüfen)')

para('Pool-Modell & Pakete – eine Vereinslizenz deckt die gesamte Frauenabteilung '
     'ab; der Hardware-Umfang skaliert mit dem gewählten Paket:', bold=True)

# Package table
pkg = doc.add_table(rows=4, cols=5)
pkg.style = 'Light List Accent 1'
hdr = ['Paket', 'Preis / Monat', 'GPS-Tracker', 'Polar H10', 'Teams']
for i, txt in enumerate(hdr):
    run = pkg.rows[0].cells[i].paragraphs[0].add_run(txt)
    run.bold = True
rows = [
    ['Starter', '219 €', '15', '15', 'bis zu 2'],
    ['Pro', '349 €', '25', '25', 'bis zu 4'],
    ['Elite', '459 €', '35', '35', 'unbegrenzt'],
]
for r_i, row in enumerate(rows, start=1):
    for c_i, val in enumerate(row):
        pkg.rows[r_i].cells[c_i].paragraphs[0].add_run(val)
doc.add_paragraph()

para('Zusätzlich sind Trageschutzwesten, Ersatzbänder (z. B. 48 Stück) sowie das '
     'Lade- und Checkout-Dock enthalten. Die Hardware wird über eine einmalige '
     'Anzahlung (2.200–4.200 €) und die laufende Lizenz finanziert; sie verbleibt im '
     'Nutzungsrecht der Abteilung. Die geteilte Nutzung erzeugt zugleich '
     'Wechselkosten: Konfiguration und individuelle Baselines der Spielerinnen '
     'entstehen über Monate und gingen bei einem Anbieterwechsel verloren.')

para('Kostenbasis & Beschaffungsannahmen – konservative, aber realistische '
     'Einkaufsannahmen für die Finanzplanung:', bold=True)
bullet('ECG-Brustgurt (HR/HRV)', 'Dual-Mode-Gurte wie Magene H64 (≈ 20–30 €), Coospo H6 '
       '(≈ 24–35 €) oder Decathlon KALENJI HRM (≈ 30–35 €). Arbeitsannahme: 30 € pro Gurt.')
bullet('GNSS-/GPS-Tracker', 'Integrierbare GNSS-Boards statt fertiger 200-€-Tracker – '
       'z. B. Seeed Wio WM1110 (≈ 30–45 €) oder Quectel L76K (≈ 19–30 €). '
       'Arbeitsannahme: 30 € pro Einheit (Board + Antenne).')
bullet('Sensor-Hardware je Spielerin', '≈ 60 € reine Sensor-Stückkosten (Brustgurt + GNSS-Board).')
bullet('Set-Kalkulation', 'Bei z. B. 20 Einheiten ≈ 1.200 € reine Sensor-Hardware; '
       'inkl. Gehäuse, Westen, Dock und Reserve rund 2.500–3.000 € pro Set – '
       'konsistent mit der einmaligen Anzahlung von 2.200–4.200 €.')

# --- 3.1.2 Software ---
h('3.1.2  Software: Decision Layer statt Rohdaten-Dashboard', 3)
para('Die Software ist das Wertschöpfungszentrum und der zentrale '
     'Differenzierungsfaktor. Während der direkte Wettbewerb vor allem Rohdaten- und '
     'Performance-Dashboards für sportwissenschaftliche Stäbe liefert, übersetzt '
     'BetterSquad die Daten in handlungsleitende Empfehlungen, die auch ohne '
     'Sportwissenschaftler:in vor Ort verständlich sind. Als Software-as-a-Service '
     'ist die Plattform für den Verein wartungsfrei und wird kontinuierlich '
     'weiterentwickelt.')

para('Trainer:innen-Dashboard & Session Planner – die Kommandozentrale für das '
     'Funktionsteam:', bold=True)
bullet('Squad-Ampel', 'Grün/Gelb/Rot-Status je Spielerin auf einen Blick')
bullet('ACWR-Chart', 'Verlauf der durchschnittlichen und maximalen Team-ACWR über 21 Tage')
bullet('KPI-Board', 'Live-KPIs zur Zahl fitter, gefährdeter und Risiko-Spielerinnen')
bullet('Spielerin-Detail', 'Detailansicht je Spielerin mit Last-, HRV- und Hooper-Historie')
bullet('Session Planner', 'Kaderverfügbarkeits-Matrix mit empfohlener individueller Intensität, exportierbar als Sessionplan')

para('Spieler:innen-App – bindet die Athletin aktiv in den Steuerungskreislauf ein:', bold=True)
bullet('Readiness', 'Tägliche Readiness-Karte mit persönlicher ACWR und Hooper-Score')
bullet('Check-in', 'Wellness-Check-in über Schlaf, Muskelkater, Energie und Stimmung')
bullet('Scan', 'QR-Equipment-Checkout mit geführter Geräteanlage')
bullet('Stats', 'Persönliche Statistiken: Distanz, Sprints, Höchstgeschwindigkeit, HRV-Trend')
bullet('Zyklus (optional)', 'Freiwillige Zyklusangabe – ausschließlich durch die Spielerin selbst')

para('Female-informed Decision Layer – die Analyse-Engine verarbeitet äußere Last, '
     'innere Last und subjektives Befinden zu einer priorisierten, individuellen '
     'Empfehlung:', bold=True)
bullet('ACWR-Berechnung', 'Acute:Chronic Workload Ratio aus 7-Tage-Akut- und '
       '28-Tage-Chroniklast; Schwellen: > 1,3 erhöht (Gelb), > 1,49 kritisch (Rot).')
bullet('Individuelle Baseline', 'Auswertung gegen eine auf weibliche Physiologie '
       'kalibrierte Baseline jeder Spielerin – nicht gegen aus Männermessungen abgeleitete Referenzwerte.')
bullet('Zykluskontext', 'Die freiwillige Zyklusangabe fließt als ein Faktor unter mehreren '
       'in den Score ein und kontextualisiert insbesondere die HRV-Interpretation.')
bullet('Klartext-Empfehlung', 'Verständliche Ampel mit Begründungstext und konkreter '
       'Handlungsempfehlung je Spielerin und fürs Gesamtteam.')

# --- 3.1.3 Datenschutz ---
h('3.1.3  Datenschutz-Architektur', 3)
para('Da BetterSquad besondere Kategorien personenbezogener Daten im Sinne des '
     'Art. 9 DSGVO (Gesundheitsdaten) verarbeitet, ist Datenschutz integraler '
     'Produktbestandteil und zugleich vertrieblicher Wettbewerbsvorteil. Die '
     'Verarbeitung erfolgt DSGVO-konform mit Serverstandort in Deutschland, '
     'verschlüsselter Übertragung und rollenbasiert getrennten Zugriffen. '
     'Entscheidend ist die bewusst gegenläufige Architektur gegenüber zyklusbasierten '
     'Plattformen: Die freiwillige Zyklusangabe ist ausschließlich ein interner '
     'Score-Faktor der Spielerin und zu keinem Zeitpunkt für Trainer:innen sichtbar. '
     'BetterSquad ist damit ausdrücklich keine Zyklus-App, sondern ein '
     'Belastungs-Tool mit Zykluskontext.')
bullet('Privacy by Design', 'Zyklusangabe freiwillig, nur durch die Spielerin, niemals als Coach-Information')
bullet('Akzeptanz', 'Akzeptanz im überwiegend männlichen Trainerstab des Amateurfußballs')
bullet('Compliance', 'Serverstandort Deutschland, Art. 9 DSGVO, verschlüsselte Übertragung')

# --- 3.1.4 Wissenschaftliche Fundierung ---
h('3.1.4  Wissenschaftliche Fundierung', 3)
para('Alle Kennzahlen des Systems beruhen auf etablierten sportwissenschaftlichen '
     'Methoden. Das grenzt BetterSquad von rein technikgetriebenen Tracking-Gadgets '
     'ab und schafft Glaubwürdigkeit gegenüber Vereinen und Verbänden:')
bullet('Gabbett (2016)', 'ACWR-Belastungssteuerung als validiertes Modell zur Verletzungsprävention.')
bullet('Saw et al. (2016)', 'Aussagekraft subjektiver Selbstauskünfte (Hooper-Index) im Wellness-Check-in.')
bullet('Plews et al. (2013)', 'HRV als anerkannter Marker für Trainingsadaptation und Regeneration.')
bullet('McNulty et al. (2020)', 'Geschlechtsspezifische Kalibrierung auf Basis der Evidenz zum Zykluseinfluss auf die Leistungsfähigkeit.')

# ============================================================
# 3.2  Unique Selling Proposition (USP)
# ============================================================
h('3.2  Unique Selling Proposition (USP)', 2)

para('BetterSquad besetzt eine strukturelle Lücke im Markt: Etablierte '
     'Hardware-Anbieter richten sich mit Pro-Spieler:in-Preislogik und überwiegend '
     'aus Männermessungen abgeleiteten Referenzwerten an Elite- und Profistrukturen, '
     'während female-spezifische Software-Plattformen zwar die Frauenphysiologie '
     'adressieren, aber weder Vereins-Hardware noch eine objektive, belastungsbasierte '
     'Steuerung bieten. BetterSquad besetzt genau die Schnittstelle beider Welten – '
     'als einziger Anbieter, der integrierte Hardware, female-informed Decision Layer '
     'und ein vereinstaugliches Pool-Preismodell verbindet.')

para('Die Alleinstellung ergibt sich aus dem Zusammenspiel von fünf Faktoren:', bold=True)
bullet('Integration', 'Äußere und innere Last sowie subjektives Befinden in einer Plattform '
       'statt getrennter Insellösungen.')
bullet('Pool-Modell', 'Eine Vereinslizenz und ein geteilter Gerätepool für die gesamte '
       'Frauenabteilung statt Pro-Gerät-Logik – Profi-Methodik zum Amateur-Preis.')
bullet('Female-informed', 'Individuelle, auf weibliche Physiologie kalibrierte Baselines '
       'mit Zyklusangabe als Score-Faktor statt männlich abgeleiteter Referenzwerte.')
bullet('Einfachheit', 'Ampel-Logik und Klartext-Empfehlungen – bedienbar ohne '
       'Sportwissenschaftler:in vor Ort.')
bullet('Datenschutz', 'DSGVO-konform, Serverstandort Deutschland, Zyklusdaten konsequent '
       'vor dem Trainerstab verborgen – Datenschutz als Vertriebsargument.')

callout('Kundennutzen der Software',
        'Aus Zahlen werden Entscheidungen: Die Plattform sagt nicht nur, dass eine '
        'Spielerin gefährdet ist, sondern auch, was zu tun ist – z. B. „Intensität um '
        '20 % reduzieren, Fokus Taktik + Regeneration". Das ersetzt die '
        'Interpretationsleistung, die im Profibereich teures Fachpersonal erbringt.')

callout('Einordnung der Kostenannahme',
        'Fertige Teamlösungen (z. B. STATSports-, Tracktics- oder '
        'Action-Tracer-Teamsets) liegen bei rund 200 € pro Tracker bzw. etwa 2.400 € '
        'für 12 Geräte. Die konservativ kalkulierten Stückkosten von BetterSquad '
        'bleiben deutlich darunter, ohne unrealistisch niedrig zu sein – die Annahmen '
        'sind damit gegenüber Vereinen, Investor:innen und Jury robust verteidigbar.')

# ============================================================
# 3.3  Entwicklungsstand und -schritte
# ============================================================
h('3.3  Entwicklungsstand und -schritte', 2)
para('Der funktionsfähige Frontend-Prototyp aus Trainer:innen-Dashboard und '
     'Spieler:innen-App bildet den vollständigen Nutzungsablauf bereits ab und dient '
     'als Grundlage für Pilotvereine und Pitch.')
bullet('Status quo', 'Validierter, klickbarer Prototyp des End-to-End-Nutzungsflusses.')
bullet('Phase 1', 'Hardware-Anbindung (Live-Daten GPS + Polar H10), Cloud-Backend, '
       'Pilot in 1–2 Frauenabteilungen.')
bullet('Phase 2', 'Schärfung des Decision Layers auf realen Saisondaten, '
       'Reporting-Export, Verbandsschnittstellen.')
bullet('Phase 3', 'Skalierung über die DACH-Landesverbände auf rund 85 Vereine bis Jahr 3.')

# ---------------- Literaturverzeichnis ----------------
h('Literaturverzeichnis', 2)
refs = [
 'Gabbett, T. J. (2016). The training—injury prevention paradox: Should athletes be '
 'training smarter and harder? British Journal of Sports Medicine, 50(5), 273–280. '
 'https://doi.org/10.1136/bjsports-2015-095788',
 'McNulty, K. L., Elliott-Sale, K. J., Dolan, E., Swinton, P. A., Ansdell, P., '
 'Goodall, S., Thomas, K., & Hicks, K. M. (2020). The effects of menstrual cycle '
 'phase on exercise performance in eumenorrheic women: A systematic review and '
 'meta-analysis. Sports Medicine, 50(10), 1813–1827. '
 'https://doi.org/10.1007/s40279-020-01319-3',
 'Plews, D. J., Laursen, P. B., Stanley, J., Kilding, A. E., & Buchheit, M. (2013). '
 'Training adaptation and heart rate variability in elite endurance athletes: '
 'Opening the door to effective monitoring. Sports Medicine, 43(9), 773–781. '
 'https://doi.org/10.1007/s40279-013-0071-8',
 'Polar Electro. (n.d.). Polar H10 heart rate sensor [Product page]. Polar Electro. '
 'https://www.polar.com/de/sensoren/h10-herzfrequenz-sensor',
 'Saw, A. E., Main, L. C., & Gastin, P. B. (2016). Monitoring the athlete training '
 'response: Subjective self-reported measures trump commonly used objective measures '
 '– A systematic review. British Journal of Sports Medicine, 50(5), 281–291. '
 'https://doi.org/10.1136/bjsports-2015-094758',
]
for ref in refs:
    doc.add_paragraph(ref, style='List Paragraph')

# Footer line
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('BetterSquad · Businessplan Teil 3: Produkt / Dienstleistung')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

out = '/home/user/bettersquad/BetterSquad_Businessplan_3_Produkt.docx'
doc.save(out)
print('saved', out)
