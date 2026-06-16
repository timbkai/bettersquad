# -*- coding: utf-8 -*-
"""Generiert Teil 3 des BetterSquad-Businessplans: Produkt / Dienstleistung (Hardware / Software).

Inhalt: an die Wettbewerbsanalyse angeglichen (Fokus Amateur-/Semi-Pro-Frauenfußball,
female-informed Decision Layer, Datenschutz-Architektur, APA-Quellen).
Format: BetterSquad-Branding, nummerierte Überschriften, Tabelle, Callout-Boxen.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Farbschema (BetterSquad Emerald / Dark) ---
EMERALD = RGBColor(0x0E, 0x96, 0x5A)
DARK = RGBColor(0x0A, 0x0F, 0x1A)
GREY = RGBColor(0x55, 0x5F, 0x6E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# --- Standard-Schrift ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(text, number=None):
    p = doc.add_heading(level=1)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run((f"{number}  " if number else "") + text)
    run.font.color.rgb = EMERALD
    run.font.size = Pt(20)
    run.font.bold = True
    return p


def h2(text, number=None):
    p = doc.add_heading(level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run((f"{number}  " if number else "") + text)
    run.font.color.rgb = DARK
    run.font.size = Pt(14)
    run.font.bold = True
    return p


def h3(text, number=None):
    p = doc.add_heading(level=3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run((f"{number}  " if number else "") + text)
    run.font.color.rgb = EMERALD
    run.font.size = Pt(12)
    run.font.bold = True
    return p


def para(text, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def rich(parts):
    """parts: Liste von (text, bold) Tupeln."""
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = EMERALD
    p.add_run(text)
    return p


def callout(title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "ECFDF5")
    cell.paragraphs[0].text = ""
    pt = cell.paragraphs[0]
    rt = pt.add_run(title)
    rt.bold = True
    rt.font.color.rgb = EMERALD
    rt.font.size = Pt(11)
    pb = cell.add_paragraph()
    rb = pb.add_run(text)
    rb.font.size = Pt(10.5)
    return table


# =====================================================================
# TITEL
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
trun = title.add_run("BetterSquad")
trun.font.size = Pt(28)
trun.font.bold = True
trun.font.color.rgb = EMERALD
sub = title.add_run("  |  Businessplan")
sub.font.size = Pt(16)
sub.font.color.rgb = GREY

para("Investiere in deinen Kader. Nicht in einzelne Spieler:innen.", italic=True, color=GREY, size=12)
para("WiSo Köln · Bachelorseminar · Wintersemester 2024/25", color=GREY, size=10)

doc.add_paragraph()

h1("Produkt / Dienstleistung", number="3")
para(
    "Dieses Kapitel beschreibt das Leistungsangebot von BetterSquad entlang der beiden "
    "untrennbaren Bausteine Hardware und Software. BetterSquad ist ein integriertes "
    "Belastungssteuerungs-System für den Amateur- und Semi-Pro-Frauenfußball im DACH-Raum, das "
    "professionelle Methodik der Belastungssteuerung für Frauenabteilungen verfügbar macht – als "
    "schlüsselfertiges Komplettsystem aus tragbarer Sensorik und cloudbasierter Analysesoftware."
)

# ---------------------------------------------------------------------
h2("Produktüberblick und Wertversprechen", number="3.1")
para(
    "BetterSquad verbindet vereinseigene Sensor-Hardware im Pool-Modell mit einer Software-"
    "Plattform, die rohe Mess- und Befindlichkeitsdaten nicht als Performance-Dashboard ausgibt, "
    "sondern zu einer konkreten, individuell auf die Spielerin kalibrierten Trainingsempfehlung "
    "verdichtet. Der inhaltliche Kern ist damit kein Tracking-Gadget, sondern ein female-informed "
    "Decision Layer: eine Entscheidungsschicht, die die äußere Belastung (GPS), die innere "
    "Belastung (Herzfrequenz und HRV) und das subjektive Befinden (Hooper-Index) auf Basis "
    "individueller Baselines jeder Spielerin auswertet."
)
para(
    "Das Angebot adressiert eine strukturelle Lücke: Etablierte Hardware-Anbieter richten sich mit "
    "Pro-Spieler:in-Preislogik und überwiegend aus Männermessungen abgeleiteten Referenzwerten an "
    "Elite- und Profistrukturen, während female-spezifische Software-Plattformen zwar die "
    "Frauenphysiologie adressieren, aber weder Vereins-Hardware noch eine objektive, "
    "belastungsbasierte Steuerung bieten. BetterSquad besetzt genau die Schnittstelle beider Welten."
)
rich([
    ("Das Produkt besteht aus drei aufeinander abgestimmten Bausteinen: ", False),
    ("(1) einer vereinseigenen Hardware ", True),
    ("im geteilten Sensor-Pool, ", False),
    ("(2) einer Software-Plattform ", True),
    ("aus Trainer:innen-Dashboard und Spieler:innen-App sowie ", False),
    ("(3) einem female-informed Decision Layer ", True),
    ("mit datenschutzfreundlicher Architektur.", False),
])
callout(
    "Kernversprechen",
    "Profi-Methodik der Belastungssteuerung zum Amateur-Preis: ab 219 €/Monat statt vier- bis "
    "fünfstelliger Jahresaufwände bei Pro-Gerät-Modellen – DSGVO-konform, mit Serverstandort "
    "Deutschland und female-informed Decision Layer.",
)

# =====================================================================
# HARDWARE
# =====================================================================
h2("Hardware: vereinseigener Sensor-Pool", number="3.2")
para(
    "Die Hardware bildet die objektive Datengrundlage des Systems. Bewusst setzt BetterSquad auf "
    "erprobte, kalibrierte Standardkomponenten statt auf teure Eigenentwicklungen. Das senkt "
    "Stückkosten und Ausfallrisiko, verkürzt die Time-to-Market und sichert eine validierte "
    "Messqualität. Die Geräte werden nicht – wie im Pro-Spieler:in-Modell der etablierten Anbieter "
    "– einzeln pro Athletin beschafft, sondern der Frauenabteilung als geteilter Pool zur Verfügung "
    "gestellt. Eine typische Abteilung mit drei bis fünf Teams und 60 bis 80 Spielerinnen wird so "
    "mit einem überschaubaren Gerätebestand abgedeckt."
)

h3("GPS-Tracker (äußere Belastung)", number="3.2.1")
para(
    "Der GPS-Tracker wird in einer leichten Trageweste zwischen den Schulterblättern getragen und "
    "erfasst die äußere Belastung in Echtzeit:"
)
bullet("Zurückgelegte Distanz und Distanz in Hochintensitäts-Zonen", bold_prefix="Laufleistung: ")
bullet("Anzahl der Sprints sowie Höchstgeschwindigkeit (km/h)", bold_prefix="Sprints & Speed: ")
bullet("High-Intensity-Runs als Indikator azyklischer Belastung", bold_prefix="HI-Runs: ")
bullet("Beschleunigungs- und Abbrems-Lasten zur Erkennung von Asymmetrien", bold_prefix="Acc/Dec: ")
para(
    "Aus diesen Rohdaten wird die externe Tageslast (Load) berechnet, die als Grundgröße in das "
    "ACWR-Modell einfließt.",
    color=GREY, size=10,
)

h3("Polar-H10-Pulsgurte (innere Belastung & Recovery)", number="3.2.2")
para(
    "Der medizinisch validierte Brustgurt Polar H10 misst die innere Belastung und liefert die "
    "Datenbasis für das Recovery-Monitoring:"
)
bullet("Herzfrequenz in Echtzeit zur Steuerung der Trainingsintensität", bold_prefix="HR: ")
bullet("Herzfrequenzvariabilität (HRV) als sensibler Frühindikator der Regeneration", bold_prefix="HRV: ")
bullet("Trendanalyse über Tage, um Überlastung mehrere Tage früher zu erkennen", bold_prefix="Recovery: ")
para(
    "Der Polar H10 ist ein etablierter Industriestandard mit hoher EKG-Genauigkeit und Bluetooth-"
    "Konnektivität – das reduziert Beschaffungs- und Wartungsaufwand erheblich.",
    color=GREY, size=10,
)

h3("Tracker-Dock & QR-Checkout", number="3.2.3")
para(
    "Die Geräte lagern und laden in einem zentralen Tracker-Dock. Vor dem Training scannt jede "
    "Spielerin über die App einen QR-Code am Dock und bekommt automatisch einen freien Tracker "
    "sowie einen Pulsgurt zugewiesen. Das löst zwei zentrale Praxisprobleme des Amateurbetriebs: "
    "Es macht eine feste 1:1-Zuordnung überflüssig (Voraussetzung des Pool-Modells) und stellt die "
    "korrekte, verwechslungsfreie Datenzuordnung sicher, ohne dass Personal Geräte manuell "
    "verteilen muss."
)
bullet("Automatische Zuweisung von Tracker- und Gurt-Nummer per QR-Scan", bold_prefix="Self-Checkout: ")
bullet("Live-Übersicht über verfügbare, ausgegebene und in Wartung befindliche Geräte", bold_prefix="Pool-Transparenz: ")
bullet("Geführte Schritt-für-Schritt-Anlage (Weste, Gurt befeuchten, Bluetooth prüfen)", bold_prefix="Onboarding: ")

h3("Pool-Modell & Pakete", number="3.2.4")
para(
    "Eine Vereinslizenz deckt die gesamte Frauenabteilung ab; der Hardware-Umfang skaliert mit dem "
    "gewählten Paket:"
)

tbl = doc.add_table(rows=4, cols=5)
tbl.style = "Light Grid Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, t in enumerate(["Paket", "Preis / Monat", "GPS-Tracker", "Polar H10", "Teams"]):
    hdr[i].paragraphs[0].text = ""
    r = hdr[i].paragraphs[0].add_run(t)
    r.bold = True
    r.font.color.rgb = WHITE
    shade_cell(hdr[i], "0E965A")
rows_data = [
    ("Starter", "219 €", "15", "15", "bis zu 2"),
    ("Pro", "349 €", "25", "25", "bis zu 4"),
    ("Elite", "459 €", "35", "35", "unbegrenzt"),
]
for ri, row in enumerate(rows_data, start=1):
    for ci, val in enumerate(row):
        cell = tbl.rows[ri].cells[ci]
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(val)
        if ci == 0:
            run.bold = True
para(
    "Zusätzlich sind Trageschutzwesten, Ersatzbänder (z. B. 48 Stück) sowie das Lade- und Checkout-"
    "Dock enthalten. Die Hardware wird über eine einmalige Anzahlung (2.200–4.200 €) und die "
    "laufende Lizenz finanziert; sie verbleibt im Nutzungsrecht der Abteilung. Die geteilte Nutzung "
    "ist nicht nur ein Preisvorteil, sondern erzeugt zugleich Wechselkosten: Konfiguration und "
    "individuelle Baselines der Spielerinnen entstehen über Monate und gingen bei einem "
    "Anbieterwechsel verloren.",
    color=GREY, size=10,
)

# =====================================================================
# SOFTWARE
# =====================================================================
h2("Software: Decision Layer statt Rohdaten-Dashboard", number="3.3")
para(
    "Die Software ist das Wertschöpfungszentrum und der zentrale Differenzierungsfaktor. Während "
    "der direkte Wettbewerb vor allem Rohdaten- und Performance-Dashboards für sportwissen"
    "schaftliche Stäbe liefert, übersetzt BetterSquad die Daten in handlungsleitende Empfehlungen, "
    "die auch ohne Sportwissenschaftler:in vor Ort verständlich sind. Als Software-as-a-Service ist "
    "die Plattform für den Verein wartungsfrei und wird kontinuierlich weiterentwickelt."
)

h3("Trainer:innen-Dashboard & Session Planner", number="3.3.1")
para("Das Web-Dashboard ist die Kommandozentrale für das Funktionsteam:")
bullet("Squad-Ampel mit Grün/Gelb/Rot-Status je Spielerin auf einen Blick", bold_prefix="Squad-Ampel: ")
bullet("Verlauf der durchschnittlichen und maximalen Team-ACWR über 21 Tage", bold_prefix="ACWR-Chart: ")
bullet("Live-KPIs zur Zahl fitter, gefährdeter und Risiko-Spielerinnen", bold_prefix="KPI-Board: ")
bullet("Detailansicht je Spielerin mit Last-, HRV- und Hooper-Historie", bold_prefix="Spielerin-Detail: ")
bullet(
    "Kaderverfügbarkeits-Matrix mit empfohlener individueller Intensität, exportierbar als Sessionplan",
    bold_prefix="Session Planner: ",
)

h3("Spieler:innen-App", number="3.3.2")
para("Die mobile App bindet die Athletin aktiv in den Steuerungskreislauf ein:")
bullet("Tägliche Readiness-Karte mit persönlicher ACWR und Hooper-Score", bold_prefix="Readiness: ")
bullet("Wellness-Check-in über Schlaf, Muskelkater, Energie und Stimmung", bold_prefix="Check-in: ")
bullet("QR-Equipment-Checkout mit geführter Geräteanlage", bold_prefix="Scan: ")
bullet("Persönliche Statistiken: Distanz, Sprints, Höchstgeschwindigkeit, HRV-Trend", bold_prefix="Stats: ")
bullet("Freiwillige Zyklusangabe – ausschließlich durch die Spielerin selbst", bold_prefix="Zyklus (optional): ")

h3("Female-informed Decision Layer", number="3.3.3")
para(
    "Der Decision Layer ist das eigentliche Alleinstellungsmerkmal der Software. Die Analyse-Engine "
    "verarbeitet äußere Last, innere Last und subjektives Befinden zu einer priorisierten, "
    "individuellen Empfehlung:"
)
bullet(
    "Acute:Chronic Workload Ratio aus 7-Tage-Akut- und 28-Tage-Chroniklast; Schwellen: > 1,3 "
    "erhöht (Gelb), > 1,49 kritisch (Rot).",
    bold_prefix="ACWR-Berechnung: ",
)
bullet(
    "Auswertung gegen eine individuelle, auf weibliche Physiologie kalibrierte Baseline jeder "
    "Spielerin – nicht gegen aus Männermessungen abgeleitete Referenzwerte.",
    bold_prefix="Individuelle Baseline: ",
)
bullet(
    "Die freiwillige Zyklusangabe fließt als ein Faktor unter mehreren in den Score ein und "
    "kontextualisiert insbesondere die HRV-Interpretation.",
    bold_prefix="Zykluskontext: ",
)
bullet(
    "Verständliche Ampel mit Begründungstext und konkreter Handlungsempfehlung je Spielerin und "
    "fürs Gesamtteam.",
    bold_prefix="Klartext-Empfehlung: ",
)
callout(
    "Kundennutzen der Software",
    "Aus Zahlen werden Entscheidungen: Die Plattform sagt nicht nur, dass eine Spielerin gefährdet "
    "ist, sondern auch, was zu tun ist – z. B. „Intensität um 20 % reduzieren, Fokus Taktik + "
    "Regeneration“. Das ersetzt die Interpretationsleistung, die im Profibereich teures "
    "Fachpersonal erbringt.",
)

# =====================================================================
# DATENSCHUTZ
# =====================================================================
h2("Datenschutz-Architektur: keine Zyklus-App, sondern ein Belastungs-Tool mit Zykluskontext", number="3.4")
para(
    "Da BetterSquad besondere Kategorien personenbezogener Daten im Sinne des Art. 9 DSGVO "
    "(Gesundheitsdaten) verarbeitet, ist Datenschutz integraler Produktbestandteil und zugleich "
    "vertrieblicher Wettbewerbsvorteil. Die gesamte Verarbeitung erfolgt DSGVO-konform mit "
    "Serverstandort in Deutschland, verschlüsselter Übertragung und rollenbasiert getrennten "
    "Zugriffen."
)
para(
    "Entscheidend ist die bewusst gegenläufige Architektur gegenüber zyklusbasierten Software-"
    "Plattformen, die Trainer:innen bei Einwilligung weitreichende Einsicht in Zyklusphase und "
    "Symptomhistorie gewähren. Bei BetterSquad ist die freiwillige Zyklusangabe ausschließlich ein "
    "interner Score-Faktor der Spielerin und zu keinem Zeitpunkt für Trainer:innen sichtbar – diese "
    "sehen nur Ampel und Begründungstext. BetterSquad ist damit ausdrücklich keine Zyklus-App, "
    "sondern ein Belastungs-Tool mit Zykluskontext."
)
bullet("Zyklusangabe freiwillig, nur durch die Spielerin, niemals als Coach-Information", bold_prefix="Privacy by Design: ")
bullet("Akzeptanz im überwiegend männlichen Trainerstab des Amateurfußballs", bold_prefix="Akzeptanz: ")
bullet("Serverstandort Deutschland, Art. 9 DSGVO, verschlüsselte Übertragung", bold_prefix="Compliance: ")

# =====================================================================
# WISSENSCHAFT
# =====================================================================
h2("Wissenschaftliche Fundierung", number="3.5")
para(
    "Alle Kennzahlen des Systems beruhen auf etablierten sportwissenschaftlichen Methoden. Das "
    "grenzt BetterSquad von rein technikgetriebenen Tracking-Gadgets ab und schafft Glaubwürdigkeit "
    "gegenüber Vereinen und Verbänden:"
)
bullet(
    "ACWR-Belastungssteuerung nach Gabbett (2016) als validiertes Modell zur Verletzungsprävention.",
    bold_prefix="Gabbett (2016): ",
)
bullet(
    "Wellness-Check-in nach der von Saw et al. (2016) bestätigten Aussagekraft subjektiver "
    "Selbstauskünfte (Hooper-Index).",
    bold_prefix="Saw et al. (2016): ",
)
bullet(
    "HRV als anerkannter Marker für Trainingsadaptation und Regeneration (Plews et al., 2013).",
    bold_prefix="Plews et al. (2013): ",
)
bullet(
    "Geschlechtsspezifische Kalibrierung auf Basis der Evidenz zum Zyklus­einfluss auf die "
    "Leistungsfähigkeit (McNulty et al., 2020).",
    bold_prefix="McNulty et al. (2020): ",
)

# =====================================================================
# KUNDENNUTZEN
# =====================================================================
h2("Kundennutzen & Alleinstellung", number="3.6")
para("Der Kundennutzen ergibt sich aus dem Zusammenspiel der Bausteine:")
bullet(
    "Äußere und innere Last sowie subjektives Befinden in einer Plattform statt getrennter "
    "Insellösungen.",
    bold_prefix="Integration: ",
)
bullet(
    "Eine Vereinslizenz und ein geteilter Gerätepool für die gesamte Frauenabteilung statt "
    "Pro-Gerät-Logik.",
    bold_prefix="Pool-Modell: ",
)
bullet(
    "Individuelle, auf weibliche Physiologie kalibrierte Baselines mit Zyklusangabe als "
    "Score-Faktor.",
    bold_prefix="Female-informed: ",
)
bullet(
    "Ampel-Logik und Klartext-Empfehlungen – bedienbar ohne Sportwissenschaftler:in vor Ort.",
    bold_prefix="Einfachheit: ",
)
bullet(
    "DSGVO-konform, Serverstandort Deutschland, Zyklusdaten konsequent vor dem Trainerstab "
    "verborgen.",
    bold_prefix="Datenschutz: ",
)

# =====================================================================
# ENTWICKLUNGSSTAND
# =====================================================================
h2("Entwicklungsstand & Roadmap", number="3.7")
para(
    "Der funktionsfähige Frontend-Prototyp aus Trainer:innen-Dashboard und Spieler:innen-App bildet "
    "den vollständigen Nutzungsablauf bereits ab und dient als Grundlage für Pilotvereine und Pitch."
)
bullet("Validierter, klickbarer Prototyp des End-to-End-Nutzungsflusses.", bold_prefix="Status quo: ")
bullet("Hardware-Anbindung (Live-Daten GPS + Polar H10), Cloud-Backend, Pilot in 1–2 Frauenabteilungen.", bold_prefix="Phase 1: ")
bullet("Schärfung des Decision Layers auf realen Saisondaten, Reporting-Export, Verbandsschnittstellen.", bold_prefix="Phase 2: ")
bullet("Skalierung über die DACH-Landesverbände auf rund 85 Vereine bis Jahr 3.", bold_prefix="Phase 3: ")

# =====================================================================
# LITERATURVERZEICHNIS
# =====================================================================
h2("Literaturverzeichnis", number="3.8")


def ref(text):
    p = doc.add_paragraph(style="List Paragraph")
    p.paragraph_format.space_after = Pt(8)
    p.add_run(text)
    return p


ref(
    "Gabbett, T. J. (2016). The training—injury prevention paradox: Should athletes be training "
    "smarter and harder? British Journal of Sports Medicine, 50(5), 273–280. "
    "https://doi.org/10.1136/bjsports-2015-095788"
)
ref(
    "McNulty, K. L., Elliott-Sale, K. J., Dolan, E., Swinton, P. A., Ansdell, P., Goodall, S., "
    "Thomas, K., & Hicks, K. M. (2020). The effects of menstrual cycle phase on exercise "
    "performance in eumenorrheic women: A systematic review and meta-analysis. Sports Medicine, "
    "50(10), 1813–1827. https://doi.org/10.1007/s40279-020-01319-3"
)
ref(
    "Plews, D. J., Laursen, P. B., Stanley, J., Kilding, A. E., & Buchheit, M. (2013). Training "
    "adaptation and heart rate variability in elite endurance athletes: Opening the door to "
    "effective monitoring. Sports Medicine, 43(9), 773–781. "
    "https://doi.org/10.1007/s40279-013-0071-8"
)
ref(
    "Polar Electro. (n.d.). Polar H10 heart rate sensor [Product page]. Polar Electro. "
    "https://www.polar.com/de/sensoren/h10-herzfrequenz-sensor"
)
ref(
    "Saw, A. E., Main, L. C., & Gastin, P. B. (2016). Monitoring the athlete training response: "
    "Subjective self-reported measures trump commonly used objective measures – A systematic "
    "review. British Journal of Sports Medicine, 50(5), 281–291. "
    "https://doi.org/10.1136/bjsports-2015-094758"
)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run("BetterSquad · Businessplan Teil 3: Produkt / Dienstleistung")
fr.italic = True
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

out = "/home/user/bettersquad/BetterSquad_Businessplan_3_Produkt.docx"
doc.save(out)
print("Gespeichert:", out)
