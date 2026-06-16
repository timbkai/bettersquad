#!/usr/bin/env python3
"""Erzeugt den fertigen BetterSquad-Businessplan als eine Word-Datei."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x0E, 0x9F, 0x6E)
DARK = RGBColor(0x0A, 0x0F, 0x1A)
GREY = RGBColor(0x55, 0x5F, 0x6B)

doc = Document()

# ---- Basis-Styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for m in (doc.sections[0],):
    m.top_margin = Cm(2.0)
    m.bottom_margin = Cm(2.0)
    m.left_margin = Cm(2.2)
    m.right_margin = Cm(2.2)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, color=None, size=9.5, align="left", white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if white:
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color:
        r.font.color.rgb = color


def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = GREEN if level == 1 else DARK
    r.font.size = Pt(14 if level == 1 else 11.5)
    return p


def body(text):
    p = doc.add_paragraph(text)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10.5)
    return p


def table(headers, rows, widths=None, total_row=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, "0E9F6E")
        set_cell(c, h, bold=True, white=True, size=9.5,
                 align="left" if i == 0 else "center")
    for ridx, row in enumerate(rows):
        cells = t.add_row().cells
        is_total = total_row and ridx == len(rows) - 1
        for i, val in enumerate(row):
            set_cell(cells[i], val, bold=is_total, size=9.5,
                     align="left" if i == 0 else "center")
            if is_total:
                shade(cells[i], "E7F5EF")
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


# =====================================================================
# TITEL
# =====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("BetterSquad")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = DARK
rr = title.add_run("  ·  Businessplan")
rr.font.size = Pt(18)
rr.font.color.rgb = GREEN
rr.bold = True

sub = doc.add_paragraph()
r = sub.add_run("Bundesliga-Methodik für den Amateurfußball – ACWR-Belastungssteuerung, "
                "HRV-Recovery-Monitoring und KI-Trainingsempfehlung als Vereinslizenz.")
r.font.size = Pt(11)
r.font.color.rgb = GREY
r.italic = True

meta = doc.add_paragraph()
r = meta.add_run("WiSo Köln · Bachelorseminar 2024/25      |      Stand: Juni 2026      |      "
                 "Preismodell: Finanzplan-Variante (kundenfinanzierte Hardware)")
r.font.size = Pt(9)
r.font.color.rgb = GREY

doc.add_paragraph()

# =====================================================================
# 1 EXECUTIVE SUMMARY
# =====================================================================
heading("1  Executive Summary")
body("BetterSquad ist die erste integrierte GPS- und Recovery-Plattform für den Amateurfußball. "
     "Wir bringen wissenschaftlich validierte Belastungssteuerung (ACWR nach Gabbett et al.), "
     "HRV-Recovery-Monitoring und KI-gestützte Trainingsempfehlungen in einen Markt, der diese "
     "Methodik bislang nur aus dem Profibereich kennt – zu 219 €/Monat statt 15.000 $/Jahr.")
body("Das Geschäftsmodell kombiniert eine wiederkehrende SaaS-Lizenz (~95 % Bruttomarge) mit einer "
     "einmaligen, kostendeckenden Hardware-Anzahlung. Die Anzahlung finanziert die Hardware vollständig "
     "vor: Der Kunde finanziert das Inventar, das Unternehmen arbeitet mit negativem Working Capital. "
     "Daraus ergibt sich eine kapitaleffiziente Wachstumsstory mit positiven Unit Economics auf jeder "
     "Paketstufe.")
body("Kernkennzahlen: Blended-LTV ≈ 15.600 € bei einem implizierten CAC von rund 1.100 € (LTV/CAC ≈ 14:1), "
     "operativer Break-even bereits im Jahr 1 (Sweat Equity + kundenfinanzierte Hardware), nachhaltige "
     "Profitabilität ab Jahr 3 bei einem Ziel von 85 Vereinen.")

# =====================================================================
# 2 PROBLEM & LÖSUNG
# =====================================================================
heading("2  Problem & Lösung")
body("Verletzungen sind im Amateurfußball der größte Kader-Killer – ohne Sportwissenschaftler vor Ort "
     "wird Belastung nach Bauchgefühl gesteuert. Profimethodik existiert, ist aber an Einzelspieler-Wearables "
     "(15.000 $/Jahr/Spieler) und Spezialpersonal gebunden und damit für Amateurvereine unbezahlbar.")
body("BetterSquad löst das mit vier Bausteinen:")
bullet("ACWR & HRV – externe (GPS) und interne (Herzfrequenz) Last wissenschaftlich kombiniert.")
bullet("KI-Trainingsempfehlung – individuelle Steuerung vor jeder Einheit, ohne Personal vor Ort.")
bullet("Verletzungsprävention – proaktive Warnung bei ACWR > 1,3; Asymmetrie-Erkennung 5–10 Tage früher.")
bullet("Pool-Modell – eine Vereinslizenz für alle Mannschaften (Männer, Frauen, Jugend inklusive).")

# =====================================================================
# 3 MARKT & ZIELGRUPPE
# =====================================================================
heading("3  Markt & Zielgruppe")
body("Zielgruppe sind ambitionierte Amateur- und semiprofessionelle Vereine (Regional-, Ober-, Verbandsliga) "
     "im DACH-Raum, die mehrere Mannschaften führen und Kader systematisch entwickeln wollen. Das Pool-Modell "
     "adressiert den Verein als Ganzes statt einzelner Spieler – größerer Wert pro Abschluss, geringere "
     "Akquisekosten je Tracker. Differenzierung: DSGVO-konform, Serverstandort Deutschland, „Made in DE\".")

# =====================================================================
# 4 GESCHÄFTSMODELL & PREISE
# =====================================================================
heading("4  Geschäftsmodell & Preise")
body("Zweistufiges Modell: monatliche SaaS-Lizenz (wiederkehrend, hochmargig) plus einmalige Hardware-Anzahlung "
     "bei Vertragsabschluss. Mindestlaufzeit 48 Monate. Die Anzahlung deckt die Hardwarekosten vollständig und "
     "erwirtschaftet zusätzlich eine positive Einmalmarge.")
table(
    ["Paket", "Abo/Monat", "Anzahlung", "Tracker / Gurte", "Mannschaften"],
    [
        ["Starter", "219 €", "4.200 €", "18 / 36", "2"],
        ["Pro (beliebteste)", "349 €", "5.900 €", "25 / 50", "4"],
        ["Elite", "459 €", "8.200 €", "35 / 70", "∞"],
    ],
    widths=[4.0, 2.8, 2.8, 3.4, 3.0],
)
body("")
p = doc.add_paragraph()
r = p.add_run("Warum diese Preise (statt niedrigerer Anzahlungen)? ")
r.bold = True
r = p.add_run("Eine günstigere Anzahlungsstaffel (2.200 / 3.200 / 4.200 €) deckt die Hardwarekosten nicht: "
              "Pro läge bei rund 0 € Marge, Elite sogar bei –225 € pro Abschluss (Verlust, der mit jedem "
              "Elite-Kunden skaliert). Die gewählte Staffel sichert auf jeder Stufe eine positive Hardware-Marge, "
              "bringt mehr Cash bei Abschluss und wirkt durch das höhere Commitment churn-senkend.")

# =====================================================================
# 5 UNIT ECONOMICS
# =====================================================================
heading("5  Unit Economics")
body("Komponentenkosten (identisch über alle Pakete): GPS-Pod 55 € · Polar H10 62 € · Band 6 € · Docking 60 €. "
     "Abo-COGS ca. 17 €/Monat (Hosting, Lizenzen) → ~95 % Bruttomarge auf der Lizenz.")
table(
    ["Paket", "HW-COGS", "Anzahlung", "HW-Marge (einmalig)", "Abo-Marge/Monat"],
    [
        ["Starter", "3.390 €", "4.200 €", "+810 €", "202 €"],
        ["Pro", "4.745 €", "5.900 €", "+1.155 €", "332 €"],
        ["Elite", "6.595 €", "8.200 €", "+1.605 €", "442 €"],
    ],
    widths=[3.4, 2.8, 2.8, 3.6, 3.4],
)
body("")
body("Blended-Werte (Mix 40 % Starter / 40 % Pro / 20 % Elite): Hardware-Marge ≈ 1.107 € einmalig, "
     "Abo-Deckungsbeitrag ≈ 302 €/Monat. Über die 48-monatige Laufzeit ergibt das einen Blended-LTV von "
     "ca. 15.600 € pro Verein – bei einem implizierten CAC von rund 1.100 € liegt die LTV/CAC-Ratio bei ≈ 14:1.")

# =====================================================================
# 6 FINANZPLANUNG & BREAK-EVEN
# =====================================================================
heading("6  Finanzplanung & Break-even")
body("Annahmen: Kundenhochlauf 15 → 45 → 85 aktive Vereine (Jahresende), Paket-Mix 40/40/20, "
     "wiederkehrender Umsatz auf Basis durchschnittlich aktiver Vereine. Jahr 1 ohne Personalkosten "
     "(Sweat Equity der Gründer); Personal ab Jahr 2. Fixkosten: IT-Setup 30.000 € (Jahr 1), laufende "
     "Technik 1.800 €/Jahr, Personal 145.000 € (Jahr 2) bzw. 199.000 € (Jahr 3).")
table(
    ["Position", "Jahr 1", "Jahr 2", "Jahr 3"],
    [
        ["Aktive Vereine (Jahresende)", "15", "45", "85"],
        ["Abo-Umsatz", "28.710 €", "114.840 €", "248.820 €"],
        ["Hardware-Umsatz (Anzahlungen)", "85.200 €", "170.400 €", "227.200 €"],
        ["Gesamtumsatz", "113.910 €", "285.240 €", "476.020 €"],
        ["Deckungsbeitrag (Abo + HW)", "43.785 €", "141.930 €", "279.840 €"],
        ["Fixkosten", "33.600 €", "146.800 €", "200.800 €"],
        ["EBIT", "+10.185 €", "−4.870 €", "+79.040 €"],
        ["EBIT kumuliert", "+10.185 €", "+5.315 €", "+84.355 €"],
    ],
    widths=[6.6, 3.0, 3.0, 3.0],
    total_row=True,
)
body("")
p = doc.add_paragraph()
r = p.add_run("Break-even-Fazit: ")
r.bold = True
r = p.add_run("Durch kundenfinanzierte Hardware (negatives Working Capital) und Sweat Equity ist das kumulierte "
              "Ergebnis bereits ab Jahr 1 positiv; die Personalrampe in Jahr 2 wird vollständig aus dem Cashflow "
              "getragen. Ab Jahr 3 wächst der wiederkehrende Lizenz-Deckungsbeitrag (235.560 €) über die Fixkosten "
              "hinaus – nachhaltige Profitabilität ohne externe Kapitalrunde.")

# =====================================================================
# 7 WETTBEWERBSVORTEIL
# =====================================================================
heading("7  Wettbewerbsvorteil & Moat")
bullet("Einziger integrierter GPS+Recovery-Anbieter für Amateurfußball – kein direkter Wettbewerber im Pool-Modell.")
bullet("Datengetriebener Lock-in: 48-Monats-Verträge und akkumulierte Vereinsdaten erhöhen Wechselkosten "
       "(Shapiro & Varian) – churn-senkend, LTV-stabilisierend.")
bullet("Kapitaleffizienz: kundenfinanzierte Hardware → kein Inventar-Funding, geringe Burn-Rate.")
bullet("Vertrauens-/Commitment-Effekt der Anzahlung (Cialdini) senkt Kündigungsneigung.")

# =====================================================================
# 8 RISIKEN
# =====================================================================
heading("8  Risiken & Gegenmaßnahmen")
bullet("Adoptionsgeschwindigkeit: Anzahlung als Einstiegshürde → durch 48-Monats-Amortisation und gleichbleibend "
       "niedrige Monatsrate entschärft; Monatsrate ist der eigentliche Kaufentscheidungstreiber.")
bullet("Hardware-Lieferkette: Zweitlieferanten für Pods/Gurte; Anzahlung deckt Stückkosten vollständig.")
bullet("Datenschutz: DSGVO-Konformität und deutscher Serverstandort als Pflicht und zugleich Verkaufsargument.")

# =====================================================================
# 9 FAZIT
# =====================================================================
heading("9  Fazit – Investorensicht")
body("BetterSquad verbindet einen unbesetzten Markt mit positiven Unit Economics auf jeder Paketstufe, "
     "kundenfinanziertem Wachstum (negatives Working Capital) und einem hochmargigen, wiederkehrenden "
     "Lizenzgeschäft. Operativer Break-even im Jahr 1, nachhaltige Profitabilität ab Jahr 3 und eine "
     "LTV/CAC-Ratio von ≈ 14:1 machen das Modell kapitaleffizient und skalierbar – wachstumskapital wirkt "
     "hier als Beschleuniger, nicht als Voraussetzung.")

doc.save("/home/user/bettersquad/BetterSquad_Businessplan.docx")
print("OK -> BetterSquad_Businessplan.docx")
